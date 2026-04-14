"""Framework Documents API + Bulk Export/Import for Hierarchy, Scales, Forms, Scoring Rules."""
import os
import uuid
from io import BytesIO
from datetime import datetime, timezone

import aiofiles
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel
from fastapi.responses import FileResponse, StreamingResponse
from openpyxl import Workbook, load_workbook
from sqlalchemy import delete, select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.permissions import require_role
from app.database import get_db
from app.dependencies import get_current_user
from app.models.framework_document import FrameworkDocument
from app.models.framework_node import FrameworkNode, NodeType
from app.models.assessment_engine import (
    AssessmentScale, AssessmentScaleLevel, AssessmentFormTemplate,
    AssessmentFormField, AggregationRule,
)
from app.models.user import User

router = APIRouter(tags=["framework-management"])

DOCS_DIR = os.environ.get("FRAMEWORK_DOCS_DIR", "/app/uploads/framework_docs")

# ============ FRAMEWORK DOCUMENTS ============

@router.post("/api/frameworks/{fw_id}/documents", status_code=201)
async def upload_document(fw_id: uuid.UUID, file: UploadFile = File(...), db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    ALLOWED = {".pdf", ".docx", ".xlsx", ".pptx", ".doc", ".xls", ".png", ".jpg", ".jpeg"}
    filename = file.filename or "unnamed"
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED:
        raise HTTPException(400, f"File type '{ext}' not allowed.")
    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(400, "File exceeds 50MB limit.")
    dir_path = os.path.join(DOCS_DIR, str(fw_id))
    os.makedirs(dir_path, exist_ok=True)
    file_path = os.path.join(dir_path, f"{uuid.uuid4()}_{filename}")
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    doc = FrameworkDocument(framework_id=fw_id, file_name=filename, file_path=file_path,
        file_size=len(content), file_type=file.content_type, uploaded_by=current_user.id)
    db.add(doc)
    await db.flush()
    return {"id": str(doc.id), "file_name": doc.file_name, "file_size": doc.file_size, "file_type": doc.file_type, "uploaded_at": doc.uploaded_at.isoformat()}

@router.get("/api/frameworks/{fw_id}/documents")
async def list_documents(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    docs = (await db.execute(select(FrameworkDocument).where(FrameworkDocument.framework_id == fw_id).order_by(FrameworkDocument.uploaded_at.desc()))).scalars().all()
    return [{"id": str(d.id), "file_name": d.file_name, "file_size": d.file_size, "file_type": d.file_type,
             "description": d.description, "uploaded_at": d.uploaded_at.isoformat()} for d in docs]

@router.get("/api/frameworks/documents/{doc_id}/download")
async def download_document(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    doc = (await db.execute(select(FrameworkDocument).where(FrameworkDocument.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Document not found")
    return FileResponse(doc.file_path, filename=doc.file_name)

@router.get("/api/frameworks/documents/{doc_id}/preview")
async def preview_document(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Return document content as HTML for in-app preview."""
    from fastapi.responses import HTMLResponse
    doc = (await db.execute(select(FrameworkDocument).where(FrameworkDocument.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Document not found")
    ext = os.path.splitext(doc.file_name)[1].lower()
    try:
        if ext == ".pdf":
            # PDF: return as file for iframe
            return FileResponse(doc.file_path, filename=doc.file_name, media_type="application/pdf")
        elif ext in (".docx", ".doc"):
            import docx as python_docx
            d = python_docx.Document(doc.file_path)
            html = "<div style='font-family:Calibri,sans-serif;max-width:800px;margin:0 auto;padding:40px;'>"
            for p in d.paragraphs:
                style = p.style.name if p.style else ""
                text = p.text.strip()
                if not text: html += "<br/>"; continue
                if "Heading 1" in style: html += f"<h1 style='color:#00338D;border-bottom:2px solid #0091DA;padding-bottom:8px;'>{text}</h1>"
                elif "Heading 2" in style: html += f"<h2 style='color:#1A1A2E;'>{text}</h2>"
                elif "Heading 3" in style: html += f"<h3 style='color:#483698;'>{text}</h3>"
                else: html += f"<p style='line-height:1.6;color:#333;'>{text}</p>"
            # Tables
            for table in d.tables:
                html += "<table style='width:100%;border-collapse:collapse;margin:16px 0;'>"
                for i, row in enumerate(table.rows):
                    html += "<tr>"
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        bg = "background:#00338D;color:white;" if i == 0 else ""
                        html += f"<{tag} style='border:1px solid #ddd;padding:8px;{bg}'>{cell.text}</{tag}>"
                    html += "</tr>"
                html += "</table>"
            html += "</div>"
            return HTMLResponse(content=html)
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(doc.file_path, data_only=True)
            html = "<div style='font-family:Calibri,sans-serif;padding:20px;'>"
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                html += f"<h2 style='color:#00338D;'>{sheet_name}</h2>"
                html += "<table style='width:100%;border-collapse:collapse;margin-bottom:24px;'>"
                for i, row in enumerate(ws.iter_rows(max_row=200, values_only=True)):
                    html += "<tr>"
                    for cell in row:
                        tag = "th" if i == 0 else "td"
                        bg = "background:#00338D;color:white;" if i == 0 else ""
                        val = str(cell) if cell is not None else ""
                        html += f"<{tag} style='border:1px solid #ddd;padding:6px 10px;font-size:13px;{bg}'>{val}</{tag}>"
                    html += "</tr>"
                html += "</table>"
            html += "</div>"
            return HTMLResponse(content=html)
        elif ext in (".png", ".jpg", ".jpeg", ".gif"):
            return FileResponse(doc.file_path, filename=doc.file_name)
        else:
            raise HTTPException(400, f"Preview not supported for {ext} files")
    except HTTPException: raise
    except Exception as e:
        raise HTTPException(500, f"Failed to generate preview: {str(e)}")


@router.delete("/api/frameworks/documents/{doc_id}", status_code=204)
async def delete_document(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    doc = (await db.execute(select(FrameworkDocument).where(FrameworkDocument.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Document not found")
    if os.path.exists(doc.file_path): os.remove(doc.file_path)
    await db.delete(doc)
    await db.flush()


class BulkDeleteDocumentsRequest(BaseModel):
    ids: list[uuid.UUID]


@router.post("/api/frameworks/{fw_id}/documents/bulk-delete")
async def bulk_delete_documents(fw_id: uuid.UUID, data: BulkDeleteDocumentsRequest, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    if not data.ids:
        raise HTTPException(400, "No document IDs provided")
    result = await db.execute(
        select(FrameworkDocument).where(FrameworkDocument.id.in_(data.ids), FrameworkDocument.framework_id == fw_id)
    )
    found = result.scalars().all()
    if not found:
        raise HTTPException(404, "None of the specified documents were found")
    deleted = 0
    for doc in found:
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        await db.delete(doc)
        deleted += 1
    await db.flush()
    return {"deleted": deleted, "requested": len(data.ids), "not_found": len(data.ids) - deleted}


# ============ HIERARCHY: DELETE ALL, EXPORT EXCEL, IMPORT EXCEL ============

@router.delete("/api/frameworks/{fw_id}/hierarchy/delete-all", status_code=204)
async def delete_all_nodes(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from sqlalchemy import update
    from app.models.assessment_engine import AssessmentResponse, AssessmentNodeScore
    from app.models.ai_assessment_log import AiAssessmentLog
    # Get all node IDs for this framework
    node_ids = [n.id for n in (await db.execute(select(FrameworkNode).where(FrameworkNode.framework_id == fw_id))).scalars().all()]
    if node_ids:
        # Nullify/delete FK references
        await db.execute(update(AssessmentResponse).where(AssessmentResponse.node_id.in_(node_ids)).values(node_id=None))
        await db.execute(delete(AssessmentNodeScore).where(AssessmentNodeScore.node_id.in_(node_ids)))
        await db.execute(delete(AiAssessmentLog).where(AiAssessmentLog.node_id.in_(node_ids)))
    # Remove self-references then delete
    await db.execute(update(FrameworkNode).where(FrameworkNode.framework_id == fw_id).values(parent_id=None))
    await db.execute(delete(FrameworkNode).where(FrameworkNode.framework_id == fw_id))
    await db.flush()

@router.get("/api/frameworks/{fw_id}/hierarchy/export-excel")
async def export_nodes_excel(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    nodes = (await db.execute(select(FrameworkNode).where(FrameworkNode.framework_id == fw_id).order_by(FrameworkNode.sort_order))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Hierarchy"
    headers = ["reference_code", "name", "name_ar", "node_type", "parent_reference_code", "is_assessable", "weight", "max_score", "sort_order", "assessment_type", "maturity_level", "evidence_type", "acceptance_criteria", "acceptance_criteria_en", "spec_references", "priority"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    node_map = {str(n.id): n.reference_code for n in nodes}
    for row_idx, n in enumerate(nodes, 2):
        parent_ref = node_map.get(str(n.parent_id), "") if n.parent_id else ""
        vals = [n.reference_code, n.name, n.name_ar, n.node_type, parent_ref, n.is_assessable, float(n.weight) if n.weight else None, float(n.max_score) if n.max_score else None, n.sort_order, n.assessment_type, n.maturity_level, n.evidence_type, n.acceptance_criteria, n.acceptance_criteria_en, n.spec_references, n.priority]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=hierarchy.xlsx"})

@router.post("/api/frameworks/{fw_id}/hierarchy/import-excel")
async def import_nodes_excel(fw_id: uuid.UUID, file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers, row))
        if r.get("reference_code"): rows.append(r)
    # Check existing nodes
    existing = (await db.execute(select(FrameworkNode.reference_code).where(FrameworkNode.framework_id == fw_id))).scalars().all()
    existing_refs = set(existing)
    new_rows = [r for r in rows if r["reference_code"] not in existing_refs]
    duplicate_rows = [r for r in rows if r["reference_code"] in existing_refs]
    # Preview mode: return what would be imported
    if preview:
        return {
            "total_in_file": len(rows),
            "new_items": [{"reference_code": r["reference_code"], "name": r.get("name", ""), "node_type": r.get("node_type", "")} for r in new_rows],
            "duplicates": [{"reference_code": r["reference_code"], "name": r.get("name", "")} for r in duplicate_rows],
            "will_import": len(new_rows),
            "will_skip": len(duplicate_rows),
        }
    # Import only new nodes
    ref_to_id = {}
    for i, r in enumerate(new_rows):
        node = FrameworkNode(framework_id=fw_id, reference_code=r["reference_code"], name=r.get("name", ""),
            name_ar=r.get("name_ar"), node_type=r.get("node_type", ""),
            is_assessable=bool(r.get("is_assessable")), weight=r.get("weight"),
            max_score=r.get("max_score"),
            sort_order=r.get("sort_order") or i, depth=0, path=f"/{uuid.uuid4()}/", is_active=True,
            assessment_type=r.get("assessment_type"),
            maturity_level=int(r["maturity_level"]) if r.get("maturity_level") is not None and str(r["maturity_level"]).strip() else None,
            evidence_type=r.get("evidence_type"), acceptance_criteria=r.get("acceptance_criteria"),
            acceptance_criteria_en=r.get("acceptance_criteria_en"), spec_references=r.get("spec_references"),
            priority=r.get("priority"))
        db.add(node)
        await db.flush()
        ref_to_id[r["reference_code"]] = node.id
    # Also load existing node IDs for parent linking
    all_nodes = (await db.execute(select(FrameworkNode).where(FrameworkNode.framework_id == fw_id))).scalars().all()
    all_ref_to_id = {n.reference_code: n.id for n in all_nodes}
    all_ref_to_depth = {n.reference_code: n.depth for n in all_nodes}
    for r in new_rows:
        parent_ref = r.get("parent_reference_code")
        if parent_ref and parent_ref in all_ref_to_id and r["reference_code"] in ref_to_id:
            node = (await db.execute(select(FrameworkNode).where(FrameworkNode.id == ref_to_id[r["reference_code"]]))).scalar_one()
            node.parent_id = all_ref_to_id[parent_ref]
            node.depth = (all_ref_to_depth.get(parent_ref, 0)) + 1
    await db.flush()
    return {"imported": len(new_rows), "skipped_duplicates": len(duplicate_rows)}


# ============ SCALES: DELETE ALL, EXPORT EXCEL, IMPORT EXCEL ============

@router.delete("/api/frameworks/{fw_id}/bulk-scales/delete-all", status_code=204)
async def delete_all_scales(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from sqlalchemy import update
    scale_ids = [s.id for s in (await db.execute(select(AssessmentScale).where(AssessmentScale.framework_id == fw_id))).scalars().all()]
    if scale_ids:
        # Nullify FK references from form templates first
        await db.execute(update(AssessmentFormTemplate).where(AssessmentFormTemplate.scale_id.in_(scale_ids)).values(scale_id=None))
        await db.execute(delete(AssessmentScaleLevel).where(AssessmentScaleLevel.scale_id.in_(scale_ids)))
        await db.execute(delete(AssessmentScale).where(AssessmentScale.framework_id == fw_id))
    await db.flush()

@router.get("/api/frameworks/{fw_id}/bulk-scales/export-excel")
async def export_scales_excel(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm import selectinload
    scales = (await db.execute(select(AssessmentScale).options(selectinload(AssessmentScale.levels)).where(AssessmentScale.framework_id == fw_id))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Scales"
    headers = ["scale_name", "scale_name_ar", "scale_type", "is_cumulative", "level_value", "level_label", "level_label_ar", "level_description", "level_description_ar", "level_color"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    row_idx = 2
    for s in scales:
        for lv in sorted(s.levels, key=lambda l: l.sort_order):
            vals = [s.name, s.name_ar, s.scale_type, s.is_cumulative, float(lv.value), lv.label, lv.label_ar, lv.description, lv.description_ar, lv.color]
            for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
            row_idx += 1
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=scales.xlsx"})

@router.post("/api/frameworks/{fw_id}/bulk-scales/import-excel")
async def import_scales_excel(fw_id: uuid.UUID, file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from decimal import Decimal
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    scale_groups: dict[str, list] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers, row))
        name = r.get("scale_name")
        if name:
            if name not in scale_groups: scale_groups[name] = {"meta": r, "levels": []}
            scale_groups[name]["levels"].append(r)
    existing_names = set((await db.execute(select(AssessmentScale.name).where(AssessmentScale.framework_id == fw_id))).scalars().all())
    new_scales = {k: v for k, v in scale_groups.items() if k not in existing_names}
    dup_scales = {k: v for k, v in scale_groups.items() if k in existing_names}
    if preview:
        return {"total_in_file": len(scale_groups), "new_items": [{"name": k, "levels": len(v["levels"])} for k, v in new_scales.items()],
                "duplicates": [{"name": k} for k in dup_scales], "will_import": len(new_scales), "will_skip": len(dup_scales)}
    created = 0
    for name, data in new_scales.items():
        meta = data["meta"]
        scale = AssessmentScale(framework_id=fw_id, name=name, name_ar=meta.get("scale_name_ar"),
            scale_type=meta.get("scale_type", "ordinal"), is_cumulative=bool(meta.get("is_cumulative")))
        db.add(scale); await db.flush()
        for i, lv in enumerate(data["levels"]):
            db.add(AssessmentScaleLevel(scale_id=scale.id, value=Decimal(str(lv.get("level_value", i))),
                label=lv.get("level_label", ""), label_ar=lv.get("level_label_ar"),
                description=lv.get("level_description"), description_ar=lv.get("level_description_ar"),
                color=lv.get("level_color"), sort_order=i))
        created += 1
    await db.flush()
    return {"imported_scales": created, "skipped_duplicates": len(dup_scales)}


# ============ FORMS: DELETE ALL, EXPORT EXCEL, IMPORT EXCEL ============

@router.delete("/api/frameworks/{fw_id}/bulk-forms/delete-all", status_code=204)
async def delete_all_forms(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from app.models.assessment_engine import AssessmentResponse
    tmpl_ids = [t.id for t in (await db.execute(select(AssessmentFormTemplate).where(AssessmentFormTemplate.framework_id == fw_id))).scalars().all()]
    if tmpl_ids:
        # Nullify FK references from responses first
        from sqlalchemy import update
        await db.execute(update(AssessmentResponse).where(AssessmentResponse.template_id.in_(tmpl_ids)).values(template_id=None))
        await db.execute(delete(AssessmentFormField).where(AssessmentFormField.template_id.in_(tmpl_ids)))
        await db.execute(delete(AssessmentFormTemplate).where(AssessmentFormTemplate.framework_id == fw_id))
    await db.flush()

@router.get("/api/frameworks/{fw_id}/bulk-forms/export-excel")
async def export_forms_excel(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm import selectinload
    tmpls = (await db.execute(select(AssessmentFormTemplate).options(selectinload(AssessmentFormTemplate.fields), selectinload(AssessmentFormTemplate.node_type), selectinload(AssessmentFormTemplate.scale)).where(AssessmentFormTemplate.framework_id == fw_id))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Forms"
    headers = ["template_name", "node_type_name", "scale_name", "field_key", "field_label", "field_label_ar", "is_required", "sort_order", "placeholder", "help_text"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    row_idx = 2
    for t in tmpls:
        for f in sorted(t.fields or [], key=lambda x: x.sort_order):
            vals = [t.name, t.node_type.name if t.node_type else "", t.scale.name if t.scale else "",
                    f.field_key, f.label, f.label_ar, f.is_required, f.sort_order, f.placeholder, f.help_text]
            for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
            row_idx += 1
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=forms.xlsx"})

@router.post("/api/frameworks/{fw_id}/bulk-forms/import-excel")
async def import_forms_excel(fw_id: uuid.UUID, file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    node_types = {nt.name: nt.id for nt in (await db.execute(select(NodeType).where(NodeType.framework_id == fw_id))).scalars().all()}
    scales = {s.name: s.id for s in (await db.execute(select(AssessmentScale).where(AssessmentScale.framework_id == fw_id))).scalars().all()}
    tmpl_groups: dict[str, list] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers, row))
        name = r.get("template_name")
        if name:
            if name not in tmpl_groups: tmpl_groups[name] = {"meta": r, "fields": []}
            tmpl_groups[name]["fields"].append(r)
    existing_names = set((await db.execute(select(AssessmentFormTemplate.name).where(AssessmentFormTemplate.framework_id == fw_id))).scalars().all())
    new_tmpls = {k: v for k, v in tmpl_groups.items() if k not in existing_names}
    dup_tmpls = {k: v for k, v in tmpl_groups.items() if k in existing_names}
    if preview:
        return {"total_in_file": len(tmpl_groups), "new_items": [{"name": k, "fields": len(v["fields"])} for k, v in new_tmpls.items()],
                "duplicates": [{"name": k} for k in dup_tmpls], "will_import": len(new_tmpls), "will_skip": len(dup_tmpls)}
    created = 0
    for name, data in new_tmpls.items():
        meta = data["meta"]
        nt_id = node_types.get(meta.get("node_type_name"))
        scale_id = scales.get(meta.get("scale_name"))
        tmpl = AssessmentFormTemplate(framework_id=fw_id, name=name, node_type_id=nt_id, scale_id=scale_id)
        db.add(tmpl); await db.flush()
        for f in data["fields"]:
            db.add(AssessmentFormField(template_id=tmpl.id, field_key=f.get("field_key", ""),
                label=f.get("field_label", ""), label_ar=f.get("field_label_ar"),
                is_required=bool(f.get("is_required")), is_visible=True,
                sort_order=f.get("sort_order") or 0, placeholder=f.get("placeholder"), help_text=f.get("help_text")))
        created += 1
    await db.flush()
    return {"imported_templates": created, "skipped_duplicates": len(dup_tmpls)}


# ============ SCORING RULES: DELETE ALL, EXPORT EXCEL, IMPORT EXCEL ============

@router.delete("/api/frameworks/{fw_id}/bulk-scoring/delete-all", status_code=204)
async def delete_all_rules(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    await db.execute(delete(AggregationRule).where(AggregationRule.framework_id == fw_id))
    await db.flush()

@router.get("/api/frameworks/{fw_id}/bulk-scoring/export-excel")
async def export_rules_excel(fw_id: uuid.UUID, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm import selectinload
    rules = (await db.execute(select(AggregationRule).options(selectinload(AggregationRule.parent_node_type), selectinload(AggregationRule.child_node_type)).where(AggregationRule.framework_id == fw_id))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Scoring Rules"
    headers = ["parent_node_type", "child_node_type", "method", "minimum_acceptable", "round_to"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, r in enumerate(rules, 2):
        vals = [r.parent_node_type.name if r.parent_node_type else "", r.child_node_type.name if r.child_node_type else "",
                r.method, float(r.minimum_acceptable) if r.minimum_acceptable else None, r.round_to]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=scoring_rules.xlsx"})

@router.post("/api/frameworks/{fw_id}/bulk-scoring/import-excel")
async def import_rules_excel(fw_id: uuid.UUID, file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from decimal import Decimal
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    node_types = {nt.name: nt.id for nt in (await db.execute(select(NodeType).where(NodeType.framework_id == fw_id))).scalars().all()}
    # Check existing rules
    from sqlalchemy.orm import selectinload
    existing_rules = (await db.execute(select(AggregationRule).options(selectinload(AggregationRule.parent_node_type), selectinload(AggregationRule.child_node_type)).where(AggregationRule.framework_id == fw_id))).scalars().all()
    existing_pairs = {(r.parent_node_type.name if r.parent_node_type else "", r.child_node_type.name if r.child_node_type else "") for r in existing_rules}
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers, row))
        if r.get("parent_node_type") and r.get("child_node_type"): rows.append(r)
    new_rows = [r for r in rows if (r["parent_node_type"], r["child_node_type"]) not in existing_pairs]
    dup_rows = [r for r in rows if (r["parent_node_type"], r["child_node_type"]) in existing_pairs]
    if preview:
        return {"total_in_file": len(rows), "new_items": [{"parent": r["parent_node_type"], "child": r["child_node_type"], "method": r.get("method")} for r in new_rows],
                "duplicates": [{"parent": r["parent_node_type"], "child": r["child_node_type"]} for r in dup_rows], "will_import": len(new_rows), "will_skip": len(dup_rows)}
    created = 0
    for r in new_rows:
        parent_id = node_types.get(r["parent_node_type"])
        child_id = node_types.get(r["child_node_type"])
        if parent_id and child_id:
            db.add(AggregationRule(framework_id=fw_id, parent_node_type_id=parent_id, child_node_type_id=child_id,
                method=r.get("method", "simple_average"),
                minimum_acceptable=Decimal(str(r["minimum_acceptable"])) if r.get("minimum_acceptable") else None,
                round_to=r.get("round_to") or 2,
                created_at=datetime.now(timezone.utc), updated_at=datetime.now(timezone.utc)))
            created += 1
    await db.flush()
    return {"imported_rules": created, "skipped_duplicates": len(dup_rows)}


# ============ ASSESSED ENTITIES: EXPORT & IMPORT ============

from app.models.assessment_engine import AssessedEntity

@router.get("/api/bulk-entities/export-excel")
async def export_entities_excel(db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from app.models.regulatory_entity import RegulatoryEntity
    from sqlalchemy.orm import selectinload
    entities = (await db.execute(
        select(AssessedEntity).options(selectinload(AssessedEntity.regulatory_entity), selectinload(AssessedEntity.regulatory_entities))
        .order_by(AssessedEntity.name)
    )).unique().scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Assessed Entities"
    headers = ["name", "name_ar", "abbreviation", "entity_type", "government_category", "sector",
               "regulatory_entities", "registration_number", "contact_person", "contact_email",
               "contact_phone", "website", "primary_color", "secondary_color", "notes", "status"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, e in enumerate(entities, 2):
        # Comma-separated regulatory entity abbreviations
        reg_abbrs = ", ".join(r.abbreviation for r in (e.regulatory_entities or []) if r.abbreviation) if hasattr(e, 'regulatory_entities') and e.regulatory_entities else ""
        if not reg_abbrs and e.regulatory_entity:
            reg_abbrs = e.regulatory_entity.abbreviation or ""
        vals = [e.name, e.name_ar, e.abbreviation, e.entity_type, e.government_category, e.sector, reg_abbrs,
                e.registration_number, e.contact_person, e.contact_email, e.contact_phone, e.website,
                e.primary_color, e.secondary_color, e.notes, e.status]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=assessed_entities.xlsx"})


@router.post("/api/bulk-entities/import-excel")
async def import_entities_excel(file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from app.models.regulatory_entity import RegulatoryEntity
    from app.models.assessment_engine import entity_regulatory_entities as ere_table
    try:
        content = await file.read()
        wb = load_workbook(BytesIO(content))
        ws = wb.active
        headers = [cell.value for cell in ws[1]]
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            r = dict(zip(headers, row))
            if r.get("name"): rows.append(r)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse Excel file: {e}")
    # Purge inactive entities whose names match import file (so they don't block re-import)
    import_names = [r["name"] for r in rows]
    if import_names:
        inactive_ids_q = await db.execute(
            select(AssessedEntity.id).where(AssessedEntity.name.in_(import_names), AssessedEntity.status != "Active")
        )
        inactive_ids = [row[0] for row in inactive_ids_q.all()]
        if inactive_ids:
            await db.execute(ere_table.delete().where(ere_table.c.entity_id.in_(inactive_ids)))
            await db.execute(delete(AssessedEntity).where(AssessedEntity.id.in_(inactive_ids)))
            await db.flush()
    # Check existing by name (only remaining active entities count as duplicates)
    existing_names = set((await db.execute(select(AssessedEntity.name))).scalars().all())
    new_rows = [r for r in rows if r["name"] not in existing_names]
    dup_rows = [r for r in rows if r["name"] in existing_names]
    if preview:
        return {"total_in_file": len(rows),
                "new_items": [{"name": r["name"], "abbreviation": r.get("abbreviation", "")} for r in new_rows],
                "duplicates": [{"name": r["name"]} for r in dup_rows],
                "will_import": len(new_rows), "will_skip": len(dup_rows)}
    # Resolve regulatory entities
    reg_entities = {re.abbreviation: re.id for re in (await db.execute(select(RegulatoryEntity))).scalars().all()}
    created = 0
    try:
        for r in new_rows:
            # Parse regulatory_entities (comma-separated) or fallback to old column
            reg_str = r.get("regulatory_entities") or r.get("regulatory_entity_abbreviation") or ""
            reg_abbr_list = [a.strip() for a in str(reg_str).split(",") if a.strip()]
            reg_id = reg_entities.get(reg_abbr_list[0]) if reg_abbr_list else None
            entity = AssessedEntity(
                name=r["name"], name_ar=r.get("name_ar"), abbreviation=r.get("abbreviation"),
                entity_type=r.get("entity_type"), government_category=r.get("government_category"),
                sector=r.get("sector"), regulatory_entity_id=reg_id,
                registration_number=r.get("registration_number"), contact_person=r.get("contact_person"),
                contact_email=r.get("contact_email"), contact_phone=r.get("contact_phone"),
                website=r.get("website"), primary_color=r.get("primary_color"), secondary_color=r.get("secondary_color"),
                notes=r.get("notes"), status=r.get("status") or "Active",
            )
            db.add(entity)
            await db.flush()
            # Create junction table entries for all regulatory entities
            for abbr in reg_abbr_list:
                rid = reg_entities.get(abbr)
                if rid:
                    await db.execute(ere_table.insert().values(entity_id=entity.id, regulatory_entity_id=rid))
            created += 1
        await db.flush()
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=400, detail=f"Import failed at row {created + 1}: {e}")
    return {"imported": created, "skipped_duplicates": len(dup_rows)}


# ============ USERS: EXPORT & IMPORT ============

from app.models.user import User as UserModel
from app.core.security import get_password_hash

@router.get("/api/bulk-users/export-excel")
async def export_users_excel(db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    users = (await db.execute(select(UserModel).order_by(UserModel.name))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Users"
    headers = ["name", "email", "role", "is_active", "assessed_entity_id"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, u in enumerate(users, 2):
        vals = [u.name, u.email, u.role, u.is_active, str(u.assessed_entity_id) if u.assessed_entity_id else ""]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=users.xlsx"})

@router.post("/api/bulk-users/import-excel")
async def import_users_excel(file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers_row = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers_row, row))
        if r.get("email"): rows.append(r)
    existing_emails = set((await db.execute(select(UserModel.email))).scalars().all())
    new_rows = [r for r in rows if r["email"] not in existing_emails]
    dup_rows = [r for r in rows if r["email"] in existing_emails]
    if preview:
        return {"total_in_file": len(rows), "new_items": [{"name": r.get("name", ""), "email": r["email"], "role": r.get("role", "")} for r in new_rows],
                "duplicates": [{"name": r.get("name", ""), "email": r["email"]} for r in dup_rows], "will_import": len(new_rows), "will_skip": len(dup_rows)}
    created = 0
    for r in new_rows:
        role = r.get("role", "kpmg_user")
        if role not in ("admin", "kpmg_user", "client"): role = "kpmg_user"
        pwd = r.get("password") or "DefaultPass123!"
        db.add(UserModel(name=r.get("name", ""), email=r["email"], hashed_password=get_password_hash(pwd),
            role=role, is_active=r.get("is_active", True) if isinstance(r.get("is_active"), bool) else True))
        created += 1
    await db.flush()
    return {"imported": created, "skipped_duplicates": len(dup_rows)}


# ============ REGULATORY ENTITIES: EXPORT & IMPORT ============

from app.models.regulatory_entity import RegulatoryEntity

@router.get("/api/bulk-regulatory-entities/export-excel")
async def export_reg_entities_excel(db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    entities = (await db.execute(select(RegulatoryEntity).order_by(RegulatoryEntity.name))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Regulatory Entities"
    headers = ["name", "name_ar", "abbreviation", "description", "website", "status"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, e in enumerate(entities, 2):
        vals = [e.name, e.name_ar, e.abbreviation, e.description, e.website, e.status]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=regulatory_entities.xlsx"})

@router.post("/api/bulk-regulatory-entities/import-excel")
async def import_reg_entities_excel(file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers_row = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers_row, row))
        if r.get("abbreviation"): rows.append(r)
    existing = set((await db.execute(select(RegulatoryEntity.abbreviation).where(RegulatoryEntity.status == "Active"))).scalars().all())
    new_rows = [r for r in rows if r["abbreviation"] not in existing]
    dup_rows = [r for r in rows if r["abbreviation"] in existing]
    if preview:
        return {"total_in_file": len(rows), "new_items": [{"name": r.get("name", ""), "abbreviation": r["abbreviation"]} for r in new_rows],
                "duplicates": [{"name": r.get("name", ""), "abbreviation": r["abbreviation"]} for r in dup_rows], "will_import": len(new_rows), "will_skip": len(dup_rows)}
    created = 0
    for r in new_rows:
        db.add(RegulatoryEntity(name=r.get("name", ""), name_ar=r.get("name_ar"), abbreviation=r["abbreviation"],
            description=r.get("description"), website=r.get("website"), status=r.get("status") or "Active"))
        created += 1
    await db.flush()
    return {"imported": created, "skipped_duplicates": len(dup_rows)}


# ============ ASSESSMENT CYCLES: EXPORT & IMPORT ============

from app.models.assessment_cycle_config import AssessmentCycleConfig
from app.models.compliance_framework import ComplianceFramework

@router.get("/api/bulk-cycles/export-excel")
async def export_cycles_excel(db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm import selectinload
    cycles = (await db.execute(select(AssessmentCycleConfig).options(selectinload(AssessmentCycleConfig.framework)).order_by(AssessmentCycleConfig.cycle_name))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Assessment Cycles"
    headers = ["cycle_name", "cycle_name_ar", "framework_abbreviation", "start_date", "end_date", "status", "description"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, c in enumerate(cycles, 2):
        fw_abbr = c.framework.abbreviation if c.framework else ""
        vals = [c.cycle_name, c.cycle_name_ar, fw_abbr,
                c.start_date.isoformat() if c.start_date else "", c.end_date.isoformat() if c.end_date else "",
                c.status, c.description]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=assessment_cycles.xlsx"})

@router.post("/api/bulk-cycles/import-excel")
async def import_cycles_excel(file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    from datetime import date as date_type
    content = await file.read()
    wb = load_workbook(BytesIO(content))
    ws = wb.active
    headers_row = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(headers_row, row))
        if r.get("cycle_name"): rows.append(r)
    existing = set((await db.execute(select(AssessmentCycleConfig.cycle_name))).scalars().all())
    new_rows = [r for r in rows if r["cycle_name"] not in existing]
    dup_rows = [r for r in rows if r["cycle_name"] in existing]
    if preview:
        return {"total_in_file": len(rows), "new_items": [{"name": r["cycle_name"], "framework": r.get("framework_abbreviation", "")} for r in new_rows],
                "duplicates": [{"name": r["cycle_name"]} for r in dup_rows], "will_import": len(new_rows), "will_skip": len(dup_rows)}
    fws = {fw.abbreviation: fw.id for fw in (await db.execute(select(ComplianceFramework))).scalars().all()}
    created = 0
    for r in new_rows:
        fw_id = fws.get(r.get("framework_abbreviation"))
        start = None
        if r.get("start_date"):
            try: start = date_type.fromisoformat(str(r["start_date"])[:10])
            except: pass
        end = None
        if r.get("end_date"):
            try: end = date_type.fromisoformat(str(r["end_date"])[:10])
            except: pass
        db.add(AssessmentCycleConfig(cycle_name=r["cycle_name"], cycle_name_ar=r.get("cycle_name_ar"),
            framework_id=fw_id, start_date=start, end_date=end,
            status=r.get("status") or "Inactive", description=r.get("description")))
        created += 1
    await db.flush()
    return {"imported": created, "skipped_duplicates": len(dup_rows)}


# ============ FRAMEWORKS: EXPORT & IMPORT ============

@router.get("/api/bulk-frameworks/export-excel")
async def export_frameworks_excel(db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm import selectinload
    fws = (await db.execute(select(ComplianceFramework).options(selectinload(ComplianceFramework.entity)).order_by(ComplianceFramework.name))).scalars().all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Frameworks"
    headers = ["name", "abbreviation", "name_ar", "description", "version", "status", "icon", "regulatory_entity_abbreviation", "requires_product_assessment"]
    for i, h in enumerate(headers, 1): ws.cell(row=1, column=i, value=h)
    for row_idx, fw in enumerate(fws, 2):
        reg_abbr = fw.entity.abbreviation if fw.entity else ""
        vals = [fw.name, fw.abbreviation, fw.name_ar, fw.description, fw.version, fw.status, fw.icon, reg_abbr, fw.requires_product_assessment]
        for i, v in enumerate(vals, 1): ws.cell(row=row_idx, column=i, value=v)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=frameworks.xlsx"})

@router.post("/api/bulk-frameworks/import-excel")
async def import_frameworks_excel(file: UploadFile = File(...), preview: bool = False, db: AsyncSession = Depends(get_db), current_user: User = Depends(require_role("admin"))):
    try:
        content = await file.read()
        wb = load_workbook(BytesIO(content))
        ws = wb.active
        headers_row = [cell.value for cell in ws[1]]
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            r = dict(zip(headers_row, row))
            if r.get("abbreviation"): rows.append(r)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse Excel file: {e}")
    existing = set((await db.execute(select(ComplianceFramework.abbreviation))).scalars().all())
    new_rows = [r for r in rows if r["abbreviation"] not in existing]
    dup_rows = [r for r in rows if r["abbreviation"] in existing]
    if preview:
        return {"total_in_file": len(rows), "new_items": [{"name": r.get("name", ""), "abbreviation": r["abbreviation"]} for r in new_rows],
                "duplicates": [{"name": r.get("name", ""), "abbreviation": r["abbreviation"]} for r in dup_rows], "will_import": len(new_rows), "will_skip": len(dup_rows)}
    reg_entities = {re.abbreviation: re.id for re in (await db.execute(select(RegulatoryEntity))).scalars().all()}
    created = 0
    try:
        for r in new_rows:
            reg_id = reg_entities.get(r.get("regulatory_entity_abbreviation")) if r.get("regulatory_entity_abbreviation") else None
            db.add(ComplianceFramework(name=r.get("name", ""), abbreviation=r["abbreviation"], name_ar=r.get("name_ar"),
                description=r.get("description"), version=r.get("version"), status=r.get("status") or "Active",
                icon=r.get("icon") or "book", entity_id=reg_id,
                requires_product_assessment=bool(r.get("requires_product_assessment"))))
            created += 1
        await db.flush()
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=400, detail=f"Import failed at row {created + 1}: {e}")
    return {"imported": created, "skipped_duplicates": len(dup_rows)}
